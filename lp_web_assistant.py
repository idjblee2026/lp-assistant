import os
import re
import json
import urllib.parse
from http.server import HTTPServer, SimpleHTTPRequestHandler
from docx import Document

PORT = 4898
BASE_DIR = r"D:\LP_Agent"
INDEX_FILE = os.path.join(BASE_DIR, "lp_database.json")
SOURCE_DIR = os.path.join(BASE_DIR, "word_reports")

def build_or_load_database():
    if os.path.exists(INDEX_FILE):
        try:
            with open(INDEX_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                if data and len(data) > 0 and "full_text" in data[0] and len(data[0]["full_text"]) > 50:
                    return data
        except Exception:
            pass

    print("419개 워드 해설 문서를 정밀 분석하여 AI 음악 비서 데이터베이스를 구축합니다...")
    if not os.path.exists(SOURCE_DIR):
        return []

    files = sorted([f for f in os.listdir(SOURCE_DIR) if f.endswith(".docx") and not f.startswith("~$")])
    database = []

    for filename in files:
        doc_path = os.path.join(SOURCE_DIR, filename)
        try:
            doc = Document(doc_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            m_loc = re.search(r"위치\s*:\s*([@NL]?\d+)", full_text)
            loc = m_loc.group(1) if m_loc else ""
            if not loc:
                m_fn = re.search(r"위치_(\d+)", filename)
                loc = m_fn.group(1) if m_fn else "미확인"

            m_album = re.search(r"앨범명\s*:\s*([^\n]+)", full_text)
            album = m_album.group(1).strip() if m_album else filename.replace(".docx", "")

            m_art = re.search(r"아티스트\s*:\s*([^\n]+)", full_text)
            artist = m_art.group(1).strip() if m_art else "미확인"

            m_genre = re.search(r"장르\s*:\s*([^\n]+)", full_text)
            genre = m_genre.group(1).strip() if m_genre else "-"

            m_year = re.search(r"발매년도\s*:\s*([^\n]+)", full_text)
            year = m_year.group(1).strip() if m_year else "-"

            # 해설 본문(상세 설명) 추출
            desc_lines = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t and not any(k in t for k in ["위치:", "앨범명:", "아티스트:", "장르:", "발매년도:", "레이블:"]):
                    desc_lines.append(t)
            desc_text = "\n".join(desc_lines) if desc_lines else full_text

            database.append({
                "filename": filename,
                "location": loc,
                "album": album,
                "artist": artist,
                "genre": genre,
                "year": year,
                "full_text": full_text,
                "desc_text": desc_text
            })
        except Exception:
            continue

    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(database, f, ensure_ascii=False, indent=2)

    return database

DB = build_or_load_database()

def search_lp(q):
    query = q.strip().lower()
    if not query:
        return []

    words = query.split()
    results = []

    for item in DB:
        score = 0
        full = item.get("full_text", "").lower()
        album = item.get("album", "").lower()
        artist = item.get("artist", "").lower()
        loc = str(item.get("location", "")).lower()

        if query == loc or f"위치_{query}" in item.get("filename", "").lower():
            score += 100

        for w in words:
            if w in loc: score += 80
            if w in album or w in artist: score += 50
            if w in full: score += 20

        if score > 0:
            # 해설 미리보기 작성
            desc = item.get("desc_text", "")
            summary = ""
            for p in desc.split("\n"):
                if len(p) > 30:
                    summary = p
                    break
            if not summary:
                summary = desc[:200]
            if len(summary) > 250:
                summary = summary[:250] + "..."

            results.append((score, {
                "location": item.get("location", "미확인"),
                "artist": item.get("artist", "미확인"),
                "album": item.get("album", "미확인"),
                "genre": item.get("genre", "-"),
                "year": item.get("year", "-"),
                "filename": item.get("filename", ""),
                "summary": summary
            }))

    results.sort(key=lambda x: x[0], reverse=True)
    return [r[1] for r in results[:10]]

HTML_PAGE = """<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>나만의 LP 서재 AI 음악 비서</title>
<style>
  body { font-family: 'Malgun Gothic', sans-serif; background: #f4f6f9; margin: 0; padding: 40px 20px; display: flex; flex-direction: column; align-items: center; }
  .container { width: 100%; max-width: 850px; background: #fff; padding: 30px; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.08); }
  h2 { text-align: center; color: #1f2937; margin-bottom: 25px; }
  .search-box { display: flex; gap: 10px; margin-bottom: 15px; }
  input[type="text"] { flex: 1; padding: 13px 16px; font-size: 16px; border: 2px solid #ddd; border-radius: 8px; outline: none; }
  input[type="text"]:focus { border-color: #2563eb; }
  button { padding: 13px 24px; font-size: 16px; background: #111827; color: #fff; border: none; border-radius: 8px; cursor: pointer; font-weight: bold; }
  button:hover { background: #374151; }
  .tags { display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 25px; justify-content: center; }
  .tag { background: #f3f4f6; border: 1px solid #e5e7eb; padding: 7px 14px; border-radius: 20px; font-size: 13px; cursor: pointer; color: #4b5563; }
  .tag:hover { background: #2563eb; color: #fff; }
  .result-item { border: 1px solid #e5e7eb; border-radius: 10px; padding: 20px; margin-bottom: 18px; background: #fafafa; border-left: 6px solid #2563eb; }
  .result-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px; }
  .result-title { font-size: 18px; font-weight: bold; color: #1e3a8a; }
  .open-btn { background: #2563eb; color: #fff; padding: 6px 14px; border-radius: 6px; text-decoration: none; font-size: 13px; font-weight: bold; cursor: pointer; border: none; }
  .open-btn:hover { background: #1d4ed8; }
  .result-meta { font-size: 13px; color: #6b7280; margin-bottom: 12px; }
  .result-desc { font-size: 14px; color: #374151; line-height: 1.7; background: #ffffff; padding: 12px 15px; border-radius: 6px; border: 1px solid #e5e7eb; white-space: pre-line; }
</style>
</head>
<body>
<div class="container">
  <h2>📀 나만의 LP 서재 AI 음악 비서</h2>
  <div class="search-box">
    <input type="text" id="query" placeholder="질문을 입력하세요 (예: 비 오는 날 감성 보컬, 아침 서정적 클래식, 108번)" onkeypress="if(event.keyCode==13) doSearch();">
    <button onclick="doSearch()">검색/추천</button>
  </div>
  <div class="tags">
    <span class="tag" onclick="quickSearch('아침 서정적 클래식')">☀️ 아침 서정적 클래식</span>
    <span class="tag" onclick="quickSearch('비 오는 날 감성 보컬')">☕ 비 오는 날 감성 보컬</span>
    <span class="tag" onclick="quickSearch('밤 피아노')">🌙 밤 피아노</span>
    <span class="tag" onclick="quickSearch('모타운 레이블')">🎷 모타운 레이블</span>
    <span class="tag" onclick="quickSearch('80년대 팝')">🎸 80년대 팝</span>
  </div>
  <div id="results"></div>
</div>

<script>
function quickSearch(text) {
  document.getElementById('query').value = text;
  doSearch();
}
function openDoc(filename) {
  fetch('/open?file=' + encodeURIComponent(filename));
}
function doSearch() {
  const q = document.getElementById('query').value.trim();
  if(!q) return;
  const resDiv = document.getElementById('results');
  resDiv.innerHTML = '<p style="text-align:center; color:#888;">음반 해설을 정밀 분석 중입니다...</p>';
  
  fetch('/search?q=' + encodeURIComponent(q))
    .then(r => r.json())
    .then(data => {
      if(!data || data.length === 0) {
        resDiv.innerHTML = '<p style="text-align:center; color:#888; margin-top:20px;">일치하는 음반을 찾지 못했습니다.</p>';
        return;
      }
      let html = '';
      data.forEach(item => {
        html += `<div class="result-item">
          <div class="result-header">
            <span class="result-title">[위치: ${item.location}] ${item.artist} - ${item.album}</span>
            <button class="open-btn" onclick="openDoc('${item.filename}')">📖 해설 원본 열기</button>
          </div>
          <div class="result-meta">장르: ${item.genre} | 발매년도: ${item.year}</div>
          <div class="result-desc"><b>💡 AI 추천 해설:</b>\n${item.summary}</div>
        </div>`;
      });
      resDiv.innerHTML = html;
    })
    .catch(e => {
      resDiv.innerHTML = '<p style="text-align:center; color:red;">검색 중 오류가 발생했습니다.</p>';
    });
}
</script>
</body>
</html>
"""

class LPRequestHandler(SimpleHTTPRequestHandler):
    def do_GET(self):
        parsed_path = urllib.parse.urlparse(self.path)
        
        if parsed_path.path == "/" or parsed_path.path == "/index.html":
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.end_headers()
            self.wfile.write(HTML_PAGE.encode("utf-8"))
            return

        if parsed_path.path == "/search":
            query_params = urllib.parse.parse_qs(parsed_path.query)
            q = query_params.get("q", [""])[0]
            matched = search_lp(q)
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            self.wfile.write(json.dumps(matched, ensure_ascii=False).encode("utf-8"))
            return

        if parsed_path.path == "/open":
            query_params = urllib.parse.parse_qs(parsed_path.query)
            fname = query_params.get("file", [""])[0]
            if fname:
                doc_file = os.path.join(SOURCE_DIR, fname)
                if os.path.exists(doc_file):
                    os.startfile(doc_file)
            self.send_response(200)
            self.end_headers()
            return
            
        return super().do_GET()

if __name__ == "__main__":
    import webbrowser
    os.chdir(BASE_DIR)
    webbrowser.open(f"http://127.0.0.1:{PORT}")
    server = HTTPServer(("127.0.0.1", PORT), LPRequestHandler)
    print(f"LP 음악비서 웹 서버 구동 중: http://127.0.0.1:{PORT}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        server.server_close()
